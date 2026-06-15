"""Render de una cotización/propuesta a Word (.docx), editable por el cliente.

Espeja la plantilla del PDF (``app.quotes.pdf``) — misma paleta CiiSA, mismas
columnas de la tabla económica y mismo orden de secciones — pero genera un .docx
con python-docx (sin dependencias de sistema) para quienes prefieren editar el
documento en Word antes de enviarlo. El documento es el limpio para el cliente:
NO incluye metadata interna (precedente usado, etc.)."""
from __future__ import annotations

import io
import os
from datetime import datetime, timezone

import fitz  # PyMuPDF — rasteriza la portada del PDF para incrustarla idéntica
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

from app.quotes.pdf import render_proposal_pdf
from app.quotes.schema import QuoteDraft

_ASSETS = os.path.join(os.path.dirname(__file__), "assets")
_ISOTIPO = os.path.join(_ASSETS, "ciisa_isotipo.png")  # isotipo chico del pie

# Paleta CiiSA (misma que el PDF), como hex sin '#' para los helpers de color.
BRAND = "001C71"  # navy dominante
INDIGO = "332CC4"  # header de tabla / banda de portada
LIME = "BFF500"  # banda de categoría
CYAN = "08A7FF"  # franja divisoria
TOTALS_BG = "E7E6E6"  # bloque de totales
BRAND_LIGHT = "EAEEF4"  # datos generales (las celdas usan los RGB de abajo para texto)

_RGB_INK = RGBColor(0x1C, 0x19, 0x17)
_RGB_WHITE = RGBColor(0xFF, 0xFF, 0xFF)
_RGB_BRAND = RGBColor(0x00, 0x1C, 0x71)
_RGB_INDIGO = RGBColor(0x33, 0x2C, 0xC4)
_RGB_MUTED = RGBColor(0x6B, 0x72, 0x80)


def _money(v: float | None, cur: str | None = None) -> str:
    if v is None:
        return "—"
    s = f"{v:,.2f}"
    return f"{cur} {s}" if cur else f"${s}"


# ── Helpers de bajo nivel (color/sombreado/bordes que python-docx no expone) ──────


def _shade(cell, hex_color: str) -> None:
    """Pinta el fondo de una celda (``w:shd``)."""
    tcpr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcpr.append(shd)


def _no_borders(table) -> None:
    """Quita los bordes de una tabla (look limpio como el PDF)."""
    tbl_pr = table._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "none")
        el.set(qn("w:sz"), "0")
        el.set(qn("w:space"), "0")
        borders.append(el)
    tbl_pr.append(borders)


def _set_cell_width(cell, width: Cm) -> None:
    cell.width = width
    tcpr = cell._tc.get_or_add_tcPr()
    tcw = OxmlElement("w:tcW")
    tcw.set(qn("w:w"), str(int(width.emu / 635)))  # EMU → twips
    tcw.set(qn("w:type"), "dxa")
    tcpr.append(tcw)


def _fill_cell(
    cell,
    text: str,
    *,
    bold: bool = False,
    size: float = 9.5,
    color: RGBColor = _RGB_INK,
    align: str = "left",
    bg: str | None = None,
    valign: str = "middle",
) -> None:
    """Escribe texto en una celda con formato (reemplaza el párrafo por defecto)."""
    if bg:
        _shade(cell, bg)
    cell.vertical_alignment = {
        "top": WD_ALIGN_VERTICAL.TOP,
        "middle": WD_ALIGN_VERTICAL.CENTER,
        "bottom": WD_ALIGN_VERTICAL.BOTTOM,
    }[valign]
    para = cell.paragraphs[0]
    para.alignment = {
        "left": WD_ALIGN_PARAGRAPH.LEFT,
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "right": WD_ALIGN_PARAGRAPH.RIGHT,
    }[align]
    para.paragraph_format.space_before = Pt(2)
    para.paragraph_format.space_after = Pt(2)
    run = para.add_run(text or "")
    run.bold = bold
    run.font.size = Pt(size)
    run.font.color.rgb = color


def _para(doc, text: str, *, bold=False, size=9.5, color=_RGB_INK, align="left",
          space_after=4, space_before=0, style=None):
    p = doc.add_paragraph(style=style)
    p.alignment = {
        "left": WD_ALIGN_PARAGRAPH.LEFT,
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "right": WD_ALIGN_PARAGRAPH.RIGHT,
    }[align]
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    if text:
        run = p.add_run(text)
        run.bold = bold
        run.font.size = Pt(size)
        run.font.color.rgb = color
    return p


def _new_doc(title: str) -> Document:
    doc = Document()
    sec = doc.sections[0]
    sec.page_width = Cm(21)
    sec.page_height = Cm(29.7)  # A4
    sec.left_margin = Cm(2)
    sec.right_margin = Cm(2)
    sec.top_margin = Cm(1.8)
    sec.bottom_margin = Cm(1.8)
    style = doc.styles["Normal"]
    style.font.name = "Arial"  # métricamente cercana a la Helvetica del PDF
    style.font.size = Pt(9.5)
    style.font.color.rgb = _RGB_INK
    cp = doc.core_properties
    cp.title = title or "Cotización"
    cp.author = "CiiSA"
    return doc


_USABLE = Cm(17)  # 21 - 2 - 2


# ── Tabla económica (compartida por cotización y propuesta) ───────────────────────


def _add_econ(doc: Document, draft: QuoteDraft) -> None:
    """Tabla económica al estilo CiiSA: header indigo, banda lima de categoría,
    franja cyan y bloque de totales gris. Las columnas No. Parte y Uni. se ocultan
    si ningún ítem las trae (igual que el PDF)."""
    has_parte = any(it.no_parte for it in draft.items)
    has_uni = any(it.unidad for it in draft.items)

    # (título, fracción de ancho) — Concepto absorbe el ancho de las opcionales ausentes.
    cols: list[tuple[str, float]] = [("Ítem", 0.08)]
    if has_parte:
        cols.append(("No. Parte", 0.14))
    concepto_idx = len(cols)
    cols.append(("Concepto", 0.0))
    if has_uni:
        cols.append(("Uni.", 0.08))
    cols.append(("Cant.", 0.09))
    cols.append(("Precio Unitario", 0.15))
    cols.append(("Precio Total", 0.15))
    concepto_w = 1.0 - sum(w for _, w in cols)
    widths = [Cm(_USABLE.cm * (w or concepto_w)) for _, w in cols]
    ncols = len(cols)

    table = doc.add_table(rows=1, cols=ncols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    _no_borders(table)

    # Header
    for j, (titulo, _) in enumerate(cols):
        align = "left" if j == concepto_idx else "center"
        _fill_cell(
            table.rows[0].cells[j], titulo, bold=True, size=8, color=_RGB_WHITE,
            align=align, bg=INDIGO,
        )

    # Banda lima con la categoría (solo sobre la columna Concepto, como el PDF real).
    if draft.categoria:
        row = table.add_row()
        for j in range(ncols):
            if j == concepto_idx:
                _fill_cell(row.cells[j], draft.categoria, bold=True, size=10,
                           color=_RGB_BRAND, align="center", bg=LIME)
            else:
                _fill_cell(row.cells[j], "")

    # Ítems
    for n, it in enumerate(draft.items, start=1):
        row = table.add_row()
        c = row.cells
        _fill_cell(c[0], str(n), bold=True, size=9, color=_RGB_INDIGO, align="center")
        k = 1
        if has_parte:
            _fill_cell(c[k], it.no_parte or "", bold=True, size=9,
                       color=_RGB_INDIGO, align="center")
            k += 1
        # Concepto: servicio en negrita + descripción gris en una segunda línea.
        cell = c[concepto_idx]
        cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(it.servicio or "—")
        r.bold = True
        r.font.size = Pt(9.5)
        if it.descripcion:
            p2 = cell.add_paragraph()
            p2.paragraph_format.space_after = Pt(2)
            r2 = p2.add_run(it.descripcion)
            r2.font.size = Pt(8)
            r2.font.color.rgb = _RGB_MUTED
        k = concepto_idx + 1
        if has_uni:
            _fill_cell(c[k], it.unidad or "", size=9.5, align="center")
            k += 1
        cant = "—" if it.cantidad is None else f"{it.cantidad:g}"
        _fill_cell(c[k], cant, size=9.5, align="center")
        _fill_cell(c[k + 1], _money(it.precio_unitario), bold=True, size=9.5, align="right")
        _fill_cell(c[k + 2], _money(it.importe), size=9.5, align="right")

    if not draft.items:
        row = table.add_row()
        _fill_cell(row.cells[concepto_idx], "Sin ítems.", size=8, color=_RGB_MUTED)

    # Aplicar anchos a cada celda de cada fila (python-docx requiere setearlo por celda).
    for row in table.rows:
        for j, w in enumerate(widths):
            _set_cell_width(row.cells[j], w)

    # Franja divisoria cyan (firma visual de la tabla CiiSA).
    stripe = doc.add_table(rows=1, cols=1)
    stripe.alignment = WD_TABLE_ALIGNMENT.CENTER
    _no_borders(stripe)
    sc = stripe.rows[0].cells[0]
    _shade(sc, CYAN)
    _set_cell_width(sc, _USABLE)
    sc.paragraphs[0].paragraph_format.space_before = Pt(0)
    sc.paragraphs[0].paragraph_format.space_after = Pt(0)
    sc.paragraphs[0].add_run(" ").font.size = Pt(3)

    doc.add_paragraph().paragraph_format.space_after = Pt(2)

    # Bloque resumen: acento indigo + Término de Pago/Validez/Moneda + totales gris.
    summary = doc.add_table(rows=3, cols=6)
    summary.alignment = WD_TABLE_ALIGNMENT.CENTER
    summary.autofit = False
    _no_borders(summary)
    sw = [Cm(0.5), Cm(3.2), Cm(5.0), Cm(17 - 0.5 - 3.2 - 5.0 - 3.0 - 3.4), Cm(3.0), Cm(3.4)]
    left = [
        ("Término de Pago:", draft.termino_pago),
        ("Validez:", draft.vigencia),
        ("Moneda:", draft.moneda),
    ]
    right = [
        ("SubTotal", _money(draft.subtotal)),
        ("IVA", _money(draft.impuestos)),
        ("Total Neto", _money(draft.total)),
    ]
    for i in range(3):
        cells = summary.rows[i].cells
        _fill_cell(cells[0], "", bg=INDIGO)  # acento navy/indigo
        _fill_cell(cells[1], left[i][0], bold=True, size=8.5)
        _fill_cell(cells[2], left[i][1] or "—", size=8.5)
        _fill_cell(cells[3], "")
        last = i == 2
        _fill_cell(cells[4], right[i][0], bold=last, size=9, align="right", bg=TOTALS_BG)
        _fill_cell(cells[5], right[i][1], bold=last, size=9, align="right", bg=TOTALS_BG)
    for row in summary.rows:
        for j, w in enumerate(sw):
            _set_cell_width(row.cells[j], w)


def _add_content(doc: Document, text: str) -> None:
    """Texto plano (con viñetas '- ' y saltos de línea) → párrafos del documento."""
    for raw in (text or "").split("\n"):
        line = raw.strip()
        if not line:
            continue
        if line[:2] in ("- ", "• ", "* "):
            _para(doc, line[2:].strip(), size=9.5, space_after=3, style="List Bullet")
        else:
            _para(doc, line, size=9.5, space_after=5)


# ── Cotización económica suelta ───────────────────────────────────────────────────


def render_quote_docx(
    draft: QuoteDraft,
    title: str,
    quote_number: str | None = None,
    company_name: str = "AIDOC",
) -> bytes:
    doc = _new_doc(title or "Cotización")
    fecha = datetime.now(timezone.utc).strftime("%d/%m/%Y")

    # ── Encabezado (banda navy) ──
    header = doc.add_table(rows=1, cols=2)
    header.alignment = WD_TABLE_ALIGNMENT.CENTER
    header.autofit = False
    _no_borders(header)
    _fill_cell(header.rows[0].cells[0], "COTIZACIÓN", bold=True, size=22,
               color=_RGB_WHITE, bg=BRAND)
    right_lines = [company_name] + ([f"N° {quote_number}"] if quote_number else []) + [fecha]
    rc = header.rows[0].cells[1]
    _shade(rc, BRAND)
    rc.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    rp = rc.paragraphs[0]
    rp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for i, ln in enumerate(right_lines):
        if i:
            rp.add_run().add_break()
        r = rp.add_run(ln)
        r.font.size = Pt(9.5)
        r.font.color.rgb = _RGB_WHITE
    _set_cell_width(header.rows[0].cells[0], Cm(_USABLE.cm * 0.6))
    _set_cell_width(rc, Cm(_USABLE.cm * 0.4))

    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    # ── Datos generales ──
    meta = doc.add_table(rows=1, cols=4)
    meta.alignment = WD_TABLE_ALIGNMENT.CENTER
    meta.autofit = False
    _no_borders(meta)
    data = [("CLIENTE", draft.cliente), ("MONEDA", draft.moneda),
            ("VIGENCIA", draft.vigencia), ("FECHA", fecha)]
    for j, (label, value) in enumerate(data):
        cell = meta.rows[0].cells[j]
        _shade(cell, BRAND_LIGHT)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(2)
        lr = p.add_run(label)
        lr.bold = True
        lr.font.size = Pt(7.5)
        lr.font.color.rgb = _RGB_MUTED
        p2 = cell.add_paragraph()
        vr = p2.add_run(value or "—")
        vr.font.size = Pt(10.5)
    mw = [Cm(_USABLE.cm * f) for f in (0.4, 0.2, 0.2, 0.2)]
    for j, w in enumerate(mw):
        _set_cell_width(meta.rows[0].cells[j], w)

    doc.add_paragraph().paragraph_format.space_after = Pt(8)

    # ── Tabla económica ──
    _add_econ(doc, draft)

    # ── Condiciones / notas ──
    if draft.condiciones:
        _para(doc, "CONDICIONES", bold=True, size=9, color=_RGB_BRAND,
              space_before=14, space_after=3)
        _para(doc, draft.condiciones, size=9)
    if draft.notas:
        _para(doc, "NOTAS", bold=True, size=9, color=_RGB_BRAND,
              space_before=10, space_after=3)
        _para(doc, draft.notas, size=9)

    _add_footer(doc.sections[0])
    return _to_bytes(doc)


# ── Propuesta completa (portada idéntica al PDF + secciones editables) ─────────────


def _render_cover_png(proposal, title: str) -> io.BytesIO:
    """Rasteriza la portada (página 1) del PDF de la propuesta a PNG de alta resolución.

    Así la portada del Word queda PIXEL-IDÉNTICA a la del PDF (fondo navy, motivo,
    logo, banda blanca con QR), que es un diseño pintado sobre canvas imposible de
    reproducir como texto editable en Word."""
    pdf_bytes = render_proposal_pdf(proposal, title)
    pdf = fitz.open(stream=pdf_bytes, filetype="pdf")
    try:
        pix = pdf[0].get_pixmap(dpi=200)
        return io.BytesIO(pix.tobytes("png"))
    finally:
        pdf.close()


def _add_floating_image(doc: Document, image, x_emu: int, y_emu: int, w, h) -> None:
    """Inserta una imagen flotante anclada a la página en (x,y) con tamaño (w,h).

    python-docx solo sabe insertar imágenes en línea; acá se reescribe el ``wp:inline``
    a un ``wp:anchor`` posicionado para lograr una imagen a sangre de página completa
    sin que el flujo del documento la empuje a otra hoja."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run()
    run.add_picture(image, width=w, height=h)
    drawing = run._r.find(qn("w:drawing"))
    inline = drawing.find(qn("wp:inline"))

    anchor = OxmlElement("wp:anchor")
    for k, v in (
        ("distT", "0"), ("distB", "0"), ("distL", "0"), ("distR", "0"),
        ("simplePos", "0"), ("relativeHeight", "0"), ("behindDoc", "0"),
        ("locked", "0"), ("layoutInCell", "1"), ("allowOverlap", "1"),
    ):
        anchor.set(k, v)
    sp = OxmlElement("wp:simplePos")
    sp.set("x", "0")
    sp.set("y", "0")
    anchor.append(sp)
    ph = OxmlElement("wp:positionH")
    ph.set("relativeFrom", "page")
    oh = OxmlElement("wp:posOffset")
    oh.text = str(int(x_emu))
    ph.append(oh)
    anchor.append(ph)
    pv = OxmlElement("wp:positionV")
    pv.set("relativeFrom", "page")
    ov = OxmlElement("wp:posOffset")
    ov.text = str(int(y_emu))
    pv.append(ov)
    anchor.append(pv)
    anchor.append(inline.find(qn("wp:extent")))
    ee = inline.find(qn("wp:effectExtent"))
    if ee is None:
        ee = OxmlElement("wp:effectExtent")
        for a in ("l", "t", "r", "b"):
            ee.set(a, "0")
    anchor.append(ee)
    anchor.append(OxmlElement("wp:wrapNone"))
    anchor.append(inline.find(qn("wp:docPr")))
    cnv = inline.find(qn("wp:cNvGraphicFramePr"))
    if cnv is not None:
        anchor.append(cnv)
    anchor.append(inline.find(qn("a:graphic")))

    drawing.remove(inline)
    drawing.append(anchor)


def render_proposal_docx(proposal, title: str) -> bytes:
    doc = _new_doc(title or "Propuesta")

    # ── Portada: imagen idéntica a la del PDF, a sangre de página completa ──
    cover = doc.sections[0]
    cover.left_margin = cover.right_margin = Cm(0)
    cover.top_margin = cover.bottom_margin = Cm(0)
    try:
        _add_floating_image(
            doc, _render_cover_png(proposal, title or "Propuesta"),
            0, 0, Cm(21), Cm(29.7),
        )
    except Exception:  # noqa: BLE001 — la portada no debe romper la descarga
        _para(doc, title or "Propuesta", bold=True, size=21, color=_RGB_BRAND)

    # ── Interior: sección nueva con márgenes normales (texto/tablas editables) ──
    interior = doc.add_section(WD_SECTION.NEW_PAGE)
    interior.page_width = Cm(21)
    interior.page_height = Cm(29.7)
    interior.left_margin = interior.right_margin = Cm(2)
    interior.top_margin = interior.bottom_margin = Cm(1.8)
    for idx, sec in enumerate(proposal.secciones):
        if idx:
            doc.add_paragraph().paragraph_format.space_after = Pt(10)
        _para(doc, sec.titulo, bold=True, size=13, color=_RGB_BRAND, space_after=4)
        _hr(doc)
        if sec.key == "economica":
            _add_econ(doc, proposal.economica)
        else:
            _add_content(doc, sec.contenido)

    _add_footer(interior)
    return _to_bytes(doc)


# ── Pie / línea / serialización ───────────────────────────────────────────────────


def _hr(doc: Document) -> None:
    """Línea horizontal navy bajo el título de sección (borde inferior de un párrafo)."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    p_pr = p._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "8")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), BRAND)
    borders.append(bottom)
    p_pr.append(borders)


def _add_footer(section) -> None:
    """Pie como el del PDF: isotipo abajo a la izquierda y 'Página N' navy en
    negrita cursiva a la derecha (sin texto de contacto)."""
    section.footer.is_linked_to_previous = False
    p = section.footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.tab_stops.add_tab_stop(_USABLE, WD_TAB_ALIGNMENT.RIGHT)
    try:
        p.add_run().add_picture(_ISOTIPO, width=Cm(1.0))
    except Exception:  # noqa: BLE001
        pass
    tr = p.add_run("\tPágina ")
    tr.bold = True
    tr.italic = True
    tr.font.size = Pt(8)
    tr.font.color.rgb = _RGB_BRAND
    fr = p.add_run()
    fr.bold = True
    fr.italic = True
    fr.font.size = Pt(8)
    fr.font.color.rgb = _RGB_BRAND
    for kind, val in (("begin", None), (None, " PAGE "), ("end", None)):
        if kind:
            fc = OxmlElement("w:fldChar")
            fc.set(qn("w:fldCharType"), kind)
            fr._r.append(fc)
        else:
            it = OxmlElement("w:instrText")
            it.set(qn("xml:space"), "preserve")
            it.text = val
            fr._r.append(it)


def _to_bytes(doc: Document) -> bytes:
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
