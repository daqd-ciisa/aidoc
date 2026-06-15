"""Unit · render de cotización/propuesta a Word (.docx).

No valida el diseño visual (eso se revisa a ojo), pero sí que el render produce un
.docx no vacío, abrible y que contenga los datos clave — y que no reviente en los
casos borde (sin ítems, sin columnas opcionales, valores None)."""
from __future__ import annotations

import io

import pytest
from docx import Document

from app.quotes.docx import render_proposal_docx, render_quote_docx
from app.quotes.proposal import ProposalDraft, ProposalSection
from app.quotes.schema import QuoteDraft, QuoteItem

pytestmark = pytest.mark.unit


def _text(blob: bytes) -> str:
    """Todo el texto del .docx (párrafos + celdas de tabla)."""
    doc = Document(io.BytesIO(blob))
    parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            parts.extend(c.text for c in row.cells)
    return "\n".join(parts)


def test_render_quote_docx_has_items_and_totals():
    draft = QuoteDraft(
        cliente="ACME", moneda="MXN", categoria="Servicios PS",
        items=[
            QuoteItem(servicio="Migración", no_parte="SRV1", unidad="Serv",
                      cantidad=2, precio_unitario=100.0, importe=200.0),
        ],
        subtotal=200.0, impuestos=32.0, total=232.0,
    )
    blob = render_quote_docx(draft, "Cotización ACME", quote_number="abc123")
    assert len(blob) > 1000
    txt = _text(blob)
    assert "ACME" in txt
    assert "Migración" in txt
    assert "SRV1" in txt  # columna No. Parte presente porque el ítem la trae
    assert "232.00" in txt  # total


def test_render_quote_docx_empty_does_not_crash():
    blob = render_quote_docx(QuoteDraft(), "")
    assert len(blob) > 1000
    Document(io.BytesIO(blob))  # abrible


def test_render_proposal_docx_includes_sections_and_economica():
    proposal = ProposalDraft(
        cliente="Globex", fecha="09 de junio de 2026",
        economica=QuoteDraft(
            cliente="Globex", moneda="MXN",
            items=[QuoteItem(servicio="Soporte", cantidad=1, precio_unitario=50.0,
                             importe=50.0)],
            subtotal=50.0, impuestos=8.0, total=58.0,
        ),
        secciones=[
            ProposalSection(key="objetivo", titulo="Objetivo",
                            contenido="Texto.\n- viñeta", fuente="generado"),
            ProposalSection(key="economica", titulo="Propuesta económica",
                            contenido="", fuente="generado"),
        ],
    )
    blob = render_proposal_docx(proposal, "Propuesta Globex")
    assert len(blob) > 1000
    # La portada es una imagen idéntica al PDF (el cliente vive ahí, no como texto);
    # el interior es nativo y editable (secciones + tabla económica).
    txt = _text(blob)
    assert "Objetivo" in txt
    assert "Soporte" in txt
    assert "58.00" in txt
    # 2 secciones: la portada (margen 0) y el interior (márgenes normales).
    assert len(Document(io.BytesIO(blob)).sections) == 2
